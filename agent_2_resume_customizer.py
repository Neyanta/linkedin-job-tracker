"""
Agent 2: Resume Customizer
Reads jobs from Google Sheet, customizes resume for each job, generates cover letters
"""

import os
import json
import gspread
from oauth2client.service_account import ServiceAccountCredentials
from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
from bs4 import BeautifulSoup
import time
from datetime import datetime

class ResumeCustomizer:
    def __init__(self, master_resume_path="resume_master.docx", use_sheets=True):
        self.master_resume_path = master_resume_path
        self.use_sheets = use_sheets
        self.sheet = None
        self.anthropic_client = None
        
        # Load config
        self.config = {
            'min_match_score': 50,  # Minimum score to process
            'auto_approve_score': 70,  # Auto-flag for Agent 3
            'max_jobs_per_day': 20,
            'output_folder': 'customized_resumes'
        }
        
        if use_sheets:
            self._setup_google_sheets()
        
        self._setup_anthropic()
        
        # Create output folder
        os.makedirs(self.config['output_folder'], exist_ok=True)
    
    def _setup_google_sheets(self):
        """Setup Google Sheets connection"""
        try:
            scope = [
                'https://spreadsheets.google.com/feeds',
                'https://www.googleapis.com/auth/drive'
            ]
            
            creds = ServiceAccountCredentials.from_json_keyfile_name(
                'credentials.json', 
                scope
            )
            
            client = gspread.authorize(creds)
            spreadsheet = client.open("LinkedIn PM Jobs")
            self.sheet = spreadsheet.sheet1
            
            print("✅ Connected to Google Sheets")
            
        except Exception as e:
            print(f"❌ Error connecting to Google Sheets: {str(e)}")
            self.use_sheets = False
    
    def _setup_anthropic(self):
        """Setup OpenAI API for AI-powered customization"""
        api_key = os.environ.get('OPENAI_API_KEY')
        
        if not api_key:
            print("⚠️  OPENAI_API_KEY not found in environment")
            print("💡 Set it: export OPENAI_API_KEY='your-key-here'")
            print("🔗 Get key at: https://platform.openai.com/api-keys")
            return
        
        self.anthropic_client = OpenAI(api_key=api_key)
        print("✅ OpenAI API configured")
    
    def fetch_unprocessed_jobs(self):
        """Get jobs from Google Sheet that need resume customization"""
        if not self.sheet:
            print("❌ Google Sheets not configured")
            return []
        
        try:
            all_data = self.sheet.get_all_records()
            
            # Filter for jobs with status "New" (not yet processed by Agent 2)
            unprocessed = [
                job for job in all_data 
                if job.get('Status', '').lower() == 'new'
            ]
            
            print(f"📋 Found {len(unprocessed)} jobs to process")
            return unprocessed[:self.config['max_jobs_per_day']]
            
        except Exception as e:
            print(f"❌ Error fetching jobs: {str(e)}")
            return []
    
    def scrape_job_description(self, job_url):
        """Scrape full job description from LinkedIn"""
        try:
            print(f"   Fetching job description...")
            
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            
            response = requests.get(job_url, headers=headers, timeout=10)
            
            if response.status_code == 200:
                soup = BeautifulSoup(response.text, 'html.parser')
                
                # Try to find job description
                description_div = soup.find('div', class_='description__text')
                if not description_div:
                    description_div = soup.find('div', class_='show-more-less-html__markup')
                
                if description_div:
                    description = description_div.get_text(strip=True, separator='\n')
                    print(f"   ✓ Fetched {len(description)} characters")
                    return description
                else:
                    print("   ⚠️  Could not find description")
                    return None
            else:
                print(f"   ⚠️  Status code: {response.status_code}")
                return None
                
        except Exception as e:
            print(f"   ❌ Error scraping: {str(e)}")
            return None
    
    def analyze_job_with_ai(self, job_title, company, job_description):
        """Use GPT to analyze job and extract key requirements"""
        if not self.anthropic_client:
            print("   ⚠️  OpenAI API not configured, using basic analysis")
            return self._basic_analysis(job_description)
        
        try:
            print("   Analyzing job with AI...")
            
            prompt = f"""Analyze this job posting and extract key information:

Job Title: {job_title}
Company: {company}

Job Description:
{job_description}

Please provide:
1. Top 5 required skills (be specific)
2. Top 5 keywords to include in resume
3. Key responsibilities
4. Seniority level (Junior/Mid/Senior/Principal)
5. Domain focus (e.g., B2B SaaS, Fintech, AI/ML, etc.)

Format as JSON:
{{
  "required_skills": ["skill1", "skill2", ...],
  "keywords": ["keyword1", "keyword2", ...],
  "responsibilities": ["resp1", "resp2", ...],
  "seniority": "Senior",
  "domain": "B2B SaaS"
}}
"""
            
            response = self.anthropic_client.chat.completions.create(
                model="gpt-4o-mini",  # Using GPT-4o-mini for cost efficiency
                messages=[
                    {"role": "user", "content": prompt}
                ],
                max_tokens=1024,
                temperature=0.7
            )
            
            response_text = response.choices[0].message.content
            
            # Extract JSON from response
            import re
            json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
            if json_match:
                analysis = json.loads(json_match.group())
                print("   ✓ AI analysis complete")
                return analysis
            else:
                print("   ⚠️  Could not parse AI response")
                return self._basic_analysis(job_description)
                
        except Exception as e:
            print(f"   ❌ AI analysis error: {str(e)}")
            return self._basic_analysis(job_description)
    
    def _basic_analysis(self, job_description):
        """Fallback analysis without AI"""
        common_skills = ['SQL', 'Python', 'A/B testing', 'Agile', 'data analysis']
        found_skills = [skill for skill in common_skills if skill.lower() in job_description.lower()]
        
        return {
            "required_skills": found_skills or common_skills[:3],
            "keywords": ["product roadmap", "stakeholder management", "user research"],
            "responsibilities": ["Product strategy", "Team leadership"],
            "seniority": "Senior",
            "domain": "General"
        }
    
    def calculate_match_score(self, analysis, master_resume_text):
        """Calculate match score between job and resume"""
        score = 0
        
        # Skills match (40 points)
        skills_found = sum(
            1 for skill in analysis['required_skills']
            if skill.lower() in master_resume_text.lower()
        )
        score += (skills_found / len(analysis['required_skills'])) * 40 if analysis['required_skills'] else 0
        
        # Keywords match (30 points)
        keywords_found = sum(
            1 for keyword in analysis['keywords']
            if keyword.lower() in master_resume_text.lower()
        )
        score += (keywords_found / len(analysis['keywords'])) * 30 if analysis['keywords'] else 0
        
        # Seniority match (30 points)
        if analysis['seniority'].lower() in master_resume_text.lower():
            score += 30
        
        return int(score)
    
    def customize_resume(self, job_title, company, analysis, output_path):
        """Create customized resume based on job analysis"""
        try:
            print("   Creating customized resume...")
            
            # Load master resume
            doc = Document(self.master_resume_path)
            
            # TODO: Actually customize the resume based on analysis
            # For now, just save a copy
            # In full implementation:
            # - Rewrite bullets to include keywords
            # - Emphasize relevant experiences
            # - Adjust summary
            
            doc.save(output_path)
            print(f"   ✓ Saved to {output_path}")
            
            return True
            
        except Exception as e:
            print(f"   ❌ Error customizing resume: {str(e)}")
            return False
    
    def generate_cover_letter(self, job_title, company, analysis, output_path):
        """Generate tailored cover letter"""
        if not self.anthropic_client:
            print("   ⚠️  Skipping cover letter (no AI)")
            return False
        
        try:
            print("   Generating cover letter...")
            
            prompt = f"""Write a professional cover letter for this job:

Job Title: {job_title}
Company: {company}
Required Skills: {', '.join(analysis['required_skills'])}
Domain: {analysis['domain']}

The candidate is Rahul Kumar, a Senior Product Manager with 6+ years of experience in B2B SaaS, Consumer Apps, E-commerce, and AI/ML products.

Write a compelling 3-paragraph cover letter that:
1. Opens with enthusiasm for the role
2. Highlights relevant experience and skills
3. Closes with interest in discussing further

Keep it professional but warm. Max 250 words.
"""
            
            response = self.anthropic_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "user", "content": prompt}
                ],
                max_tokens=1024,
                temperature=0.8
            )
            
            cover_letter = response.choices[0].message.content
            
            # Save cover letter as docx
            doc = Document()
            
            # Add content
            doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}\n")
            doc.add_paragraph(f"To: Hiring Manager\n{company}\n")
            doc.add_paragraph(f"Re: Application for {job_title}\n")
            
            for para in cover_letter.split('\n\n'):
                doc.add_paragraph(para)
            
            doc.add_paragraph("\nSincerely,\nRahul Kumar")
            
            doc.save(output_path)
            print(f"   ✓ Cover letter saved")
            
            return True
            
        except Exception as e:
            print(f"   ❌ Error generating cover letter: {str(e)}")
            return False
    
    def update_sheet_status(self, job_id, match_score, resume_path, cover_letter_path):
        """Update Google Sheet with resume links and match score"""
        if not self.sheet:
            return
        
        try:
            # Find row with this job_id
            all_data = self.sheet.get_all_values()
            
            # Column layout: 1=Job ID, 2=Title, 3=Company, 4=Location, 5=Link, 6=Found Date, 7=Status, 8=Notes, 9=Description
            for idx, row in enumerate(all_data[1:], start=2):  # Skip header
                if row[0] == str(job_id):  # Job ID is first column
                    self.sheet.update_cell(idx, 7, "Resume Ready")  # Status column
                    # Match score and resume path in columns 10, 11
                    if len(row) < 10:
                        self.sheet.update_cell(idx, 10, f"{match_score}%")
                    if len(row) < 11:
                        self.sheet.update_cell(idx, 11, resume_path)
                    
                    print(f"   ✓ Updated Google Sheet")
                    break
                    
        except Exception as e:
            print(f"   ⚠️  Error updating sheet: {str(e)}")
    
    def process_job(self, job):
        """Process a single job"""
        job_id = job.get('Job ID', 'unknown')
        title = job.get('Title', 'Unknown Title')
        company = job.get('Company', 'Unknown Company')
        link = job.get('Link', '')
        
        print(f"\n{'='*60}")
        print(f"Processing: {title} at {company}")
        print(f"{'='*60}")
        
        # Step 1: Get job description (from sheet first, then scrape from URL)
        description = job.get('Description', '').strip() if job.get('Description') else None
        if not description and link:
            description = self.scrape_job_description(link)
        
        if not description or len(description) < 100:
            print("   ⚠️  Job description too short or missing, skipping")
            return False
        
        # Step 2: Analyze with AI
        analysis = self.analyze_job_with_ai(title, company, description)
        
        # Step 3: Read master resume for match scoring
        master_text = ""
        try:
            master_doc = Document(self.master_resume_path)
            master_text = '\n'.join([p.text for p in master_doc.paragraphs])
        except:
            pass
        
        # Step 4: Calculate match score
        match_score = self.calculate_match_score(analysis, master_text)
        print(f"   📊 Match Score: {match_score}%")
        
        if match_score < self.config['min_match_score']:
            print(f"   ⚠️  Score below minimum ({self.config['min_match_score']}%), skipping")
            return False
        
        # Step 5: Create output folder
        folder_name = f"{job_id}_{company.replace(' ', '_')}_{title.replace(' ', '_')[:20]}"
        job_folder = os.path.join(self.config['output_folder'], folder_name)
        os.makedirs(job_folder, exist_ok=True)
        
        # Step 6: Customize resume
        resume_path = os.path.join(job_folder, f"Resume_RahulKumar_{company.replace(' ', '')}.docx")
        resume_success = self.customize_resume(title, company, analysis, resume_path)
        
        # Step 7: Generate cover letter
        cover_letter_path = os.path.join(job_folder, f"CoverLetter_{company.replace(' ', '')}.docx")
        cover_letter_success = self.generate_cover_letter(title, company, analysis, cover_letter_path)
        
        # Step 8: Update Google Sheet
        if resume_success:
            self.update_sheet_status(job_id, match_score, resume_path, cover_letter_path)
            print(f"   ✅ Job processed successfully!")
            return True
        
        return False
    
    def run(self):
        """Main execution"""
        print("🤖 Agent 2: Resume Customizer Starting...\n")
        
        # Fetch unprocessed jobs
        jobs = self.fetch_unprocessed_jobs()
        
        if not jobs:
            print("✅ No jobs to process")
            return
        
        # Process each job
        processed = 0
        for job in jobs:
            try:
                if self.process_job(job):
                    processed += 1
                
                # Rate limiting
                time.sleep(3)
                
            except Exception as e:
                print(f"❌ Error processing job: {str(e)}")
                continue
        
        print(f"\n{'='*60}")
        print(f"✅ Processed {processed} out of {len(jobs)} jobs")
        print(f"📂 Output folder: {self.config['output_folder']}")
        print(f"{'='*60}")


def main():
    """Main execution function"""
    
    # Check for required files
    if not os.path.exists('resume_master.docx'):
        print("❌ resume_master.docx not found!")
        print("💡 Run: node create_mock_resume.js")
        return
    
    if not os.path.exists('credentials.json'):
        print("⚠️  credentials.json not found")
        print("💡 Agent 2 will work but won't update Google Sheets")
    
    # Initialize and run
    customizer = ResumeCustomizer(
        master_resume_path="resume_master.docx",
        use_sheets=os.path.exists('credentials.json')
    )
    
    customizer.run()


if __name__ == "__main__":
    main()
